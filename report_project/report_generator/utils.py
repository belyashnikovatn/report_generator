from docx import Document
from django.conf import settings


def generate_docx_report(data, total):
    config = settings.REPORT_CONFIG
    display_fields = config['fields_to_display']

    doc = Document()

    # Заголовок
    doc.add_heading('1.6.1. Автомобильный транспорт', level=1)
    doc.add_heading(
        '1.6.1.1. Автомобильные дороги федерального, регионального и межмуниципального значения', level=2)

    doc.add_paragraph(
        'Внешние транспортные связи населенных пунктов обеспечиваются автомобильным транспортом, '
        'выполняющим основные грузовые и пассажирские перевозки. Перечень автомобильных дорог общего пользования '
        'федерального, регионального или межмуниципального значения, расположенных на территории муниципального округа, '
        'приведен в таблицах ниже.\n\n'
        'На территории муниципального округа не расположены автомобильные дороги федерального значения.'
    )

    doc.add_paragraph(
        'Таблица 1.6.1.1.1 Дороги общего пользования регионального и межмуниципального значения:')

    # Таблица
    table = doc.add_table(rows=1, cols=len(display_fields) + 1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    for i, field in enumerate(display_fields):
        hdr_cells[i + 1].text = field

    # Заполнение таблицы данными
    for idx, row in enumerate(data, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        for i, field in enumerate(display_fields):
            row_cells[i + 1].text = str(row.get(field, ''))

    # Итоговая строка с суммой под нужным столбцом
    last_column_index = len(display_fields)
    row_cells = table.add_row().cells
    row_cells[0].text = 'Итого:'
    for i in range(1, len(display_fields)):
        row_cells[i].text = ''
    row_cells[last_column_index].text = str(total)

    # Подпись под итогом
    doc.add_paragraph(
        f"Протяженность автомобильных дорог регионального или межмуниципального значения составляет {total} км."
    )

    return doc
