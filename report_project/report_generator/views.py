import csv
from io import TextIOWrapper, BytesIO
from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings
from .forms import CSVUploadForm
from docx import Document


def process_csv(csv_file):
    """Обработка CSV с учётом только строк, где REQUIRED_FIELD не пусто"""
    config = getattr(settings, 'CSV_REPORT_CONFIG', {})
    selected_fields = config.get('SELECTED_FIELDS', [])
    sum_fields = config.get('SUM_FIELDS', [])
    required_field = config.get('REQUIRED_FIELD')
    preview_rows = config.get('PREVIEW_ROWS', 5)
    
    data = []
    sums = {field: 0 for field in sum_fields}
    total_processed = 0
    total_skipped = 0
    
    csv_file = TextIOWrapper(csv_file.file, encoding='utf-8')
    reader = csv.DictReader(csv_file)
    
    for row in reader:
        # Проверяем, что обязательное поле заполнено
        required_value = row.get(required_field, '').strip() if required_field else None
        
        if required_field and not required_value:
            total_skipped += 1
            continue
            
        total_processed += 1
        
        # Подготавливаем данные для отчёта
        filtered_row = {}
        for field in selected_fields:
            value = row.get(field, '').strip()
            filtered_row[field] = value if value else '-'  # Заменяем пустые на '-'
        
        data.append(filtered_row)
        
        # Считаем суммы только для непустых числовых значений
        for field in sum_fields:
            try:
                value = row.get(field, '').strip()
                if value:  # Только если значение не пустое
                    sums[field] += float(value)
            except (ValueError, TypeError):
                continue
    
    return {
        'total_processed': total_processed,
        'total_skipped': total_skipped,
        'sample_data': data[:preview_rows],
        'columns': selected_fields,
        'sums': sums,
        'report_title': config.get('REPORT_TITLE', 'Отчёт'),
        'required_field': required_field
    }


def generate_docx(report_data):
    """Генерация DOCX файла с отчетом"""
    document = Document()

    # Заголовок отчета
    document.add_heading(report_data['report_title'], 0)
    document.add_paragraph(f'Всего обработано строк: {report_data["total_rows"]}')

    # Таблица с данными
    table = document.add_table(rows=1, cols=len(report_data["columns"]))
    hdr_cells = table.rows[0].cells

    # Заголовки столбцов
    for i, column in enumerate(report_data["columns"]):
        hdr_cells[i].text = column

    # Данные таблицы
    for row in report_data["sample_data"]:
        row_cells = table.add_row().cells
        for i, column in enumerate(report_data["columns"]):
            row_cells[i].text = str(row.get(column, ''))

    # Блок с суммами
    if report_data['sums']:
        document.add_paragraph('\nСуммы:')
        for field, total in report_data['sums'].items():
            document.add_paragraph(f'{field}: {total:.2f}')

    # Сохранение в буфер
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer


def upload_csv(request):
    if request.method == 'POST':
        form = CSVUploadForm(request.POST, request.FILES)
        if form.is_valid():
            csv_file = request.FILES['csv_file']
            report_data = process_csv(csv_file)

            # Обработка скачивания DOCX
            if 'download' in request.POST:
                buffer = generate_docx(report_data)
                response = HttpResponse(
                    buffer.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = 'attachment; filename=report.docx'
                return response

            # Отображение отчета на странице
            return render(request, 'report.html', {
                'form': form,
                'report_data': report_data,
                'show_report': True
            })

    # GET запрос - показать форму
    form = CSVUploadForm()
    return render(request, 'report.html', {
        'form': form,
        'show_report': False
    })
